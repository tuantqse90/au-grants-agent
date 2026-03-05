"""Export proposals to Markdown, DOCX, and PDF formats."""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from typing import Optional

from rich.console import Console

from au_grants_agent.config import settings
from au_grants_agent.models import Grant, Proposal
from au_grants_agent.utils.logger import get_logger

logger = get_logger()
console = Console()


def _sanitize_filename(text: str, max_len: int = 40) -> str:
    """Create a filesystem-safe short name from text."""
    text = re.sub(r"[^\w\s-]", "", text)
    text = re.sub(r"\s+", "_", text).strip("_")
    return text[:max_len]


def _get_output_dir(base_dir: Optional[Path] = None) -> Path:
    """Get or create the date-based output directory."""
    base = base_dir or settings.proposals_dir
    date_dir = base / datetime.utcnow().strftime("%Y-%m-%d")
    date_dir.mkdir(parents=True, exist_ok=True)
    return date_dir


def _build_filename(grant: Grant, ext: str) -> str:
    """Build standard filename: {go_id}_{short_title}_{date}.{ext}"""
    go_part = grant.go_id or grant.id[:8]
    title_part = _sanitize_filename(grant.title)
    date_part = datetime.utcnow().strftime("%Y%m%d")
    return f"{go_part}_{title_part}_{date_part}.{ext}"


def _find_unicode_font() -> Optional[str]:
    """Find a TTF font that supports Unicode (Vietnamese) on this system."""
    candidates = [
        # macOS
        "/Library/Fonts/Arial Unicode.ttf",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
        # Linux
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/freefont/FreeSans.ttf",
        "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
        "/usr/share/fonts/TTF/DejaVuSans.ttf",
        # Windows
        "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/arialuni.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            logger.debug("Found Unicode font: %s", path)
            return path
    return None


def _find_unicode_font_bold() -> Optional[str]:
    """Find a bold variant of a Unicode TTF font."""
    candidates = [
        "/Library/Fonts/Arial Unicode.ttf",  # macOS — same file, no bold variant
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        "/usr/share/fonts/truetype/freefont/FreeSansBold.ttf",
        "/usr/share/fonts/truetype/noto/NotoSans-Bold.ttf",
        "C:/Windows/Fonts/arialbd.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return path
    return None


class ProposalExporter:
    """Export generated proposals to multiple formats."""

    def __init__(self, output_dir: Optional[Path] = None) -> None:
        self.output_dir = output_dir

    def _full_content(self, proposal: Proposal) -> str:
        """Combine English content and Vietnamese summary."""
        parts = []
        if proposal.content_en:
            parts.append(proposal.content_en)
        if proposal.summary_vi:
            parts.append("---\n\n## Tóm tắt Tiếng Việt\n\n" + proposal.summary_vi)
        return "\n\n".join(parts)

    # ── Markdown ────────────────────────────────────────────────

    def export_markdown(self, grant: Grant, proposal: Proposal) -> Path:
        """Export proposal as Markdown with YAML frontmatter."""
        out_dir = _get_output_dir(self.output_dir)
        filename = _build_filename(grant, "md")
        filepath = out_dir / filename

        frontmatter = f"""---
title: "{grant.title}"
grant_id: "{grant.id}"
go_id: "{grant.go_id or 'N/A'}"
agency: "{grant.agency or 'N/A'}"
organisation: "{proposal.org_name or 'N/A'}"
generated: "{proposal.generated_at}"
model: "{proposal.model}"
---

"""
        filepath.write_text(frontmatter + self._full_content(proposal), encoding="utf-8")
        logger.info("Exported MD: %s", filepath)
        return filepath

    # ── DOCX ────────────────────────────────────────────────────

    def export_docx(self, grant: Grant, proposal: Proposal) -> Path:
        """Export proposal as DOCX with professional formatting."""
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        out_dir = _get_output_dir(self.output_dir)
        filename = _build_filename(grant, "docx")
        filepath = out_dir / filename

        doc = Document()

        # Default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(12)

        # Title page
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(grant.title)
        title_run.bold = True
        title_run.font.size = Pt(24)

        doc.add_paragraph()  # spacer

        meta_lines = [
            f"Grant Program: {grant.agency or 'Australian Government'}",
            f"Applicant Organisation: {proposal.org_name or 'N/A'}",
            f"Date: {datetime.utcnow().strftime('%d %B %Y')}",
            f"Grant ID: {grant.go_id or grant.id[:8]}",
        ]
        for line in meta_lines:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

        # Content — parse markdown-like sections
        content = self._full_content(proposal)
        for line in content.split("\n"):
            line = line.rstrip()
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("---"):
                doc.add_page_break()
            elif line.startswith("|"):
                doc.add_paragraph(line, style="Normal")
            elif line.strip() == "":
                doc.add_paragraph()
            else:
                p = doc.add_paragraph()
                parts = re.split(r"(\*\*.*?\*\*)", line)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)

        # Page numbers
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        section = doc.sections[-1]
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        p._element.append(fld)

        doc.save(str(filepath))
        logger.info("Exported DOCX: %s", filepath)
        return filepath

    # ── PDF ──────────────────────────────────────────────────────

    def _parse_table_rows(self, lines: list[str], start_idx: int) -> tuple[list[list[str]], int]:
        """Parse consecutive markdown table lines into a list of rows (list of cells).

        Returns (rows, end_idx) where end_idx is the index after the last table line.
        """
        rows: list[list[str]] = []
        i = start_idx
        while i < len(lines) and lines[i].strip().startswith("|"):
            line = lines[i].strip()
            # Split cells, strip outer pipes
            cells = [c.strip() for c in line.strip("|").split("|")]
            # Skip separator rows (----, :---, etc.)
            if cells and not all(
                all(ch in "-=: " for ch in cell) for cell in cells
            ):
                # Strip bold markers from cells
                cells = [re.sub(r"\*\*(.*?)\*\*", r"\1", c) for c in cells]
                rows.append(cells)
            i += 1
        return rows, i

    def _render_pdf_table(self, pdf, rows: list[list[str]], set_font_fn, safe_write_fn):
        """Render a markdown table as a proper PDF table with borders."""
        if not rows:
            return

        num_cols = max(len(r) for r in rows)
        # Calculate available width
        available_w = pdf.w - pdf.l_margin - pdf.r_margin
        col_w = available_w / num_cols

        # Limit column width for readability
        if col_w > 60:
            col_w = 60

        # Check if first row looks like a header (all non-empty)
        is_header = True
        row_h = 6

        for row_idx, row in enumerate(rows):
            # Pad row to num_cols
            while len(row) < num_cols:
                row.append("")

            # Check page space — need at least row height
            if pdf.get_y() > 265:
                pdf.add_page()

            pdf.set_x(pdf.l_margin)

            if row_idx == 0:
                # Header row
                set_font_fn("B", 8)
                pdf.set_fill_color(0, 200, 100)
                pdf.set_text_color(255, 255, 255)
                for cell in row:
                    pdf.cell(col_w, row_h, cell[:40], border=1, fill=True)
                pdf.ln(row_h)
                pdf.set_text_color(0, 0, 0)
            else:
                # Data row — alternate background
                set_font_fn("", 8)
                if row_idx % 2 == 0:
                    pdf.set_fill_color(240, 240, 240)
                    fill = True
                else:
                    fill = False
                for cell in row:
                    pdf.cell(col_w, row_h, cell[:40], border=1, fill=fill)
                pdf.ln(row_h)

        pdf.ln(3)

    def export_pdf(self, grant: Grant, proposal: Proposal) -> Path:
        """Export proposal as PDF with full Unicode support (Vietnamese)."""
        from fpdf import FPDF

        out_dir = _get_output_dir(self.output_dir)
        filename = _build_filename(grant, "pdf")
        filepath = out_dir / filename

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=25)

        # ── Load Unicode font ──
        font_path = _find_unicode_font()
        font_bold_path = _find_unicode_font_bold()
        uni_font = None

        if font_path:
            try:
                pdf.add_font("UniFont", "", font_path, uni=True)
                if font_bold_path:
                    pdf.add_font("UniFont", "B", font_bold_path, uni=True)
                else:
                    pdf.add_font("UniFont", "B", font_path, uni=True)
                uni_font = "UniFont"
                logger.info("PDF: Using Unicode font from %s", font_path)
            except Exception as e:
                logger.warning("PDF: Failed to load Unicode font: %s", e)
                uni_font = None

        def set_font(style: str = "", size: int = 11):
            if uni_font:
                pdf.set_font(uni_font, style, size)
            else:
                pdf.set_font("Helvetica", style, size)

        def safe_write(text: str, h: int = 6):
            pdf.set_x(pdf.l_margin)
            try:
                pdf.multi_cell(0, h, text)
            except UnicodeEncodeError:
                safe = text.encode("latin-1", errors="replace").decode("latin-1")
                pdf.multi_cell(0, h, safe)
            except Exception:
                pdf.ln(h)

        # ── Title Page ──
        pdf.add_page()
        pdf.ln(50)

        # Green accent bar
        pdf.set_fill_color(0, 255, 136)
        pdf.rect(15, pdf.get_y(), 3, 30, "F")

        pdf.set_x(25)
        set_font("B", 24)
        pdf.multi_cell(0, 10, grant.title)
        pdf.ln(20)

        set_font("", 12)
        meta = [
            ("Agency", grant.agency or "Australian Government"),
            ("Organisation", proposal.org_name or "N/A"),
            ("Date", datetime.utcnow().strftime("%d %B %Y")),
            ("Grant ID", grant.go_id or grant.id[:8]),
        ]
        for label, value in meta:
            pdf.set_x(pdf.l_margin)
            set_font("B", 11)
            pdf.cell(35, 8, f"{label}:")
            set_font("", 11)
            pdf.cell(0, 8, value, ln=True)

        # Bottom accent line
        pdf.ln(15)
        pdf.set_draw_color(0, 255, 136)
        pdf.set_line_width(1.0)
        pdf.line(15, pdf.get_y(), 195, pdf.get_y())

        # Footer on title page
        pdf.set_y(-30)
        set_font("", 9)
        pdf.set_text_color(120, 120, 120)
        pdf.cell(0, 6, "Generated by AU Grants Agent | NullShift", align="C", ln=True)
        pdf.set_text_color(0, 0, 0)

        # ── Content Pages ──
        pdf.add_page()
        content = self._full_content(proposal)
        lines = content.split("\n")
        toc_entries: list[tuple[str, int, int]] = []  # (title, level, page)
        i = 0

        while i < len(lines):
            line = lines[i].rstrip()

            # Check page space
            if pdf.get_y() > 265:
                pdf.add_page()

            # ── Tables (consume block) ──
            if line.strip().startswith("|"):
                table_rows, i = self._parse_table_rows(lines, i)
                if table_rows:
                    self._render_pdf_table(pdf, table_rows, set_font, safe_write)
                continue

            # ── Headings ──
            if line.startswith("### "):
                heading = line[4:]
                pdf.ln(4)
                # Thin accent line before h3
                pdf.set_draw_color(0, 200, 100)
                pdf.set_line_width(0.3)
                pdf.line(pdf.l_margin, pdf.get_y(), pdf.l_margin + 30, pdf.get_y())
                pdf.ln(2)
                set_font("B", 12)
                safe_write(heading)
                pdf.ln(2)
                toc_entries.append((heading, 3, pdf.page))
            elif line.startswith("## "):
                heading = line[3:]
                pdf.ln(6)
                set_font("B", 14)
                pdf.set_text_color(0, 150, 80)
                safe_write(heading)
                pdf.set_text_color(0, 0, 0)
                pdf.ln(3)
                toc_entries.append((heading, 2, pdf.page))
            elif line.startswith("# "):
                heading = line[2:]
                pdf.ln(8)
                # Green accent bar before h1
                pdf.set_fill_color(0, 255, 136)
                pdf.rect(pdf.l_margin, pdf.get_y(), 3, 12, "F")
                pdf.set_x(pdf.l_margin + 6)
                set_font("B", 16)
                pdf.multi_cell(0, 8, heading)
                pdf.ln(4)
                toc_entries.append((heading, 1, pdf.page))

            # ── Bold standalone line ──
            elif line.startswith("**") and line.endswith("**"):
                pdf.ln(2)
                set_font("B", 11)
                safe_write(line.strip("*"))
                set_font("", 11)

            # ── Horizontal rule ──
            elif line.startswith("---"):
                pdf.ln(6)
                pdf.set_draw_color(0, 255, 136)
                pdf.set_line_width(0.6)
                pdf.line(15, pdf.get_y(), 195, pdf.get_y())
                pdf.ln(6)

            # ── Numbered list (e.g., "1. ", "2. ") ──
            elif re.match(r"^\d+\.\s+", line):
                set_font("", 11)
                clean = re.sub(r"\*\*(.*?)\*\*", r"\1", line)
                safe_write(f"  {clean}")

            # ── Bullet point ──
            elif line.startswith("*   ") or line.startswith("- ") or line.startswith("* "):
                set_font("", 11)
                bullet_text = line.lstrip("*- ").strip()
                bullet_text = re.sub(r"\*\*(.*?)\*\*", r"\1", bullet_text)
                safe_write(f"    \u2022  {bullet_text}")

            # ── Sub-bullet / indented ──
            elif line.startswith("    "):
                set_font("", 10)
                clean = re.sub(r"\*\*(.*?)\*\*", r"\1", line.strip().lstrip("*- "))
                safe_write(f"        \u2013  {clean}")

            # ── Empty line ──
            elif line.strip() == "":
                pdf.ln(3)

            # ── Regular paragraph ──
            else:
                set_font("", 11)
                clean = re.sub(r"\*\*(.*?)\*\*", r"\1", line)
                safe_write(clean)

            i += 1

        # ── Page Headers & Footers ──
        total_pages = pdf.pages_count
        short_title = grant.title[:60] + ("..." if len(grant.title) > 60 else "")
        for pg in range(1, total_pages + 1):
            pdf.page = pg

            # Header (skip title page)
            if pg > 1:
                pdf.set_y(5)
                set_font("", 7)
                pdf.set_text_color(140, 140, 140)
                pdf.cell(0, 5, short_title, align="L")
                pdf.cell(0, 5, grant.agency or "Australian Government", align="R", ln=True)
                pdf.set_draw_color(200, 200, 200)
                pdf.set_line_width(0.2)
                pdf.line(pdf.l_margin, 11, pdf.w - pdf.r_margin, 11)
                pdf.set_text_color(0, 0, 0)

            # Footer
            pdf.set_y(-15)
            set_font("", 8)
            pdf.set_text_color(140, 140, 140)
            pdf.cell(0, 10, f"Page {pg} of {total_pages}", align="C")
            pdf.set_text_color(0, 0, 0)

        pdf.output(str(filepath))
        logger.info("Exported PDF: %s", filepath)
        console.print(f"  [#00ff88]PDF saved: {filepath}[/#00ff88]")
        return filepath

    # ── Export All ───────────────────────────────────────────────

    def export_all(
        self, grant: Grant, proposal: Proposal, formats: str = "all"
    ) -> list[Path]:
        """Export proposal in specified formats. Returns list of created file paths."""
        paths = []
        fmt = formats.lower()

        if fmt in ("all", "md"):
            paths.append(self.export_markdown(grant, proposal))

        if fmt in ("all", "docx"):
            try:
                paths.append(self.export_docx(grant, proposal))
            except Exception as e:
                logger.error("DOCX export failed: %s", e)
                console.print(f"[red]DOCX export failed: {e}[/red]")

        if fmt in ("all", "pdf"):
            try:
                paths.append(self.export_pdf(grant, proposal))
            except Exception as e:
                logger.error("PDF export failed: %s", e)
                console.print(f"[red]PDF export failed: {e}[/red]")

        return paths
